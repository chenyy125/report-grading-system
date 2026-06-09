import os
import sys
from docx import Document
from docx.shared import Pt

def create_test_docx(file_path, content_dict, has_images=False):
    doc = Document()
    
    for section_title, section_content in content_dict.items():
        heading = doc.add_heading(section_title, level=1)
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
        
        for paragraph in section_content:
            doc.add_paragraph(paragraph)
    
    if has_images:
        doc.add_picture = lambda x: None
    
    doc.save(file_path)
    print(f"Created test file: {file_path}")

def create_test_data():
    reference_content = {
        '实训目的': [
            '本实训旨在让学生掌握RAG问答系统的基本原理和实现方法。',
            '通过实践操作，理解检索增强生成技术的核心概念。',
            '学会使用向量数据库存储和检索文档内容。'
        ],
        '实训步骤': [
            '步骤一：环境准备。安装必要的Python依赖包，包括langchain、chromadb等。',
            '步骤二：数据加载。将文档内容加载到系统中，进行文本分割处理。',
            '步骤三：向量存储。使用embedding模型将文本转换为向量并存储到向量数据库。',
            '步骤四：检索问答。实现基于检索的问答功能，将用户问题与知识库进行匹配。',
            '步骤五：结果评估。测试问答系统的准确性和响应速度。'
        ],
        '实验结果': [
            '成功实现了RAG问答系统，能够准确回答知识库中的问题。',
            '系统响应时间在1秒以内，满足实时问答需求。',
            '在测试集上的准确率达到90%以上。'
        ],
        '问题反思': [
            '在实现过程中遇到了向量数据库性能优化的问题。',
            '通过调整chunk大小和使用索引，最终解决了性能瓶颈。',
            '未来可以考虑使用更高效的向量数据库如Pinecone。'
        ],
        '心得体会': [
            '通过本次实训，我深刻理解了RAG技术的优势和应用场景。',
            '学会了如何将理论知识应用到实际项目中。',
            '团队协作能力和问题解决能力得到了提升。'
        ]
    }
    
    student_good_content = {
        '实训目的': [
            '本实训旨在学习RAG问答系统的开发方法。',
            '理解检索增强生成的工作原理。'
        ],
        '实训步骤': [
            '安装Python依赖包，配置开发环境。',
            '加载文档并进行文本分割。',
            '使用embedding模型生成向量并存储到向量数据库。',
            '实现问答功能，测试系统性能。'
        ],
        '实验结果': [
            '成功实现RAG问答系统，响应速度较快。',
            '测试准确率达到85%左右。'
        ],
        '问题反思': [
            '遇到了一些技术问题，通过查阅资料解决了。',
            '对向量数据库的使用有了更深入的理解。'
        ],
        '心得体会': [
            '收获很大，学到了很多新知识。',
            '对AI应用开发有了更全面的认识。'
        ]
    }
    
    student_bad_content = {
        '实训目的': [
            '做RAG项目。'
        ],
        '实训步骤': [
            '安装软件。',
            '写代码。',
            '运行程序。'
        ],
        '实验结果': [
            '程序能运行。'
        ]
    }
    
    reference_dir = os.path.join(os.path.dirname(__file__), 'reference')
    submissions_dir = os.path.join(os.path.dirname(__file__), 'submissions')
    
    os.makedirs(reference_dir, exist_ok=True)
    os.makedirs(submissions_dir, exist_ok=True)
    
    create_test_docx(os.path.join(reference_dir, 'reference_report.docx'), reference_content)
    create_test_docx(os.path.join(submissions_dir, 'student_good.docx'), student_good_content)
    create_test_docx(os.path.join(submissions_dir, 'student_bad.docx'), student_bad_content)
    
    print("\n测试数据创建完成！")
    print(f"参考答案: {os.path.join(reference_dir, 'reference_report.docx')}")
    print(f"优秀学生报告: {os.path.join(submissions_dir, 'student_good.docx')}")
    print(f"较差学生报告: {os.path.join(submissions_dir, 'student_bad.docx')}")

def test_system():
    sys.path.insert(0, os.path.dirname(__file__))
    
    from document_parser import DocumentParser
    from comparator import ReportComparator
    from scorer import ReportScorer
    from feedback_generator import FeedbackGenerator, ReportExporter
    
    reference_path = os.path.join(os.path.dirname(__file__), 'reference', 'reference_report.docx')
    student_good_path = os.path.join(os.path.dirname(__file__), 'submissions', 'student_good.docx')
    student_bad_path = os.path.join(os.path.dirname(__file__), 'submissions', 'student_bad.docx')
    
    if not os.path.exists(reference_path) or not os.path.exists(student_good_path) or not os.path.exists(student_bad_path):
        print("测试数据不存在，正在创建...")
        create_test_data()
    
    parser = DocumentParser()
    comparator = ReportComparator()
    scorer = ReportScorer()
    feedback_generator = FeedbackGenerator()
    exporter = ReportExporter(os.path.join(os.path.dirname(__file__), 'output'))
    
    print("\n=== 测试参考答案解析 ===")
    ref_data = parser.parse_document(reference_path)
    ref_data['sections'] = parser.extract_sections(ref_data, ['实训目的', '实训步骤', '实验结果', '问题反思', '心得体会'])
    print(f"章节数量: {len(ref_data['sections'])}")
    print(f"章节名称: {list(ref_data['sections'].keys())}")
    print(f"图片数量: {ref_data['image_count']}")
    print(f"文本长度: {len(ref_data['total_text'])}")
    
    print("\n=== 测试优秀学生报告 ===")
    good_data = parser.parse_document(student_good_path)
    good_data['sections'] = parser.extract_sections(good_data, ['实训目的', '实训步骤', '实验结果', '问题反思', '心得体会'])
    
    comparison_good = comparator.compare_documents(ref_data, good_data)
    scores_good = scorer.calculate_scores(comparison_good, good_data)
    feedback_good = feedback_generator.generate_feedback(scores_good, comparison_good, good_data)
    
    print(f"总分: {scores_good['total']}/100")
    print(f"等级: {scorer.get_grade(scores_good['total'])}")
    print(f"优点: {feedback_good['strengths'][:2]}")
    print(f"不足: {feedback_good['weaknesses'][:2]}")
    
    print("\n=== 测试较差学生报告 ===")
    bad_data = parser.parse_document(student_bad_path)
    bad_data['sections'] = parser.extract_sections(bad_data, ['实训目的', '实训步骤', '实验结果', '问题反思', '心得体会'])
    
    comparison_bad = comparator.compare_documents(ref_data, bad_data)
    scores_bad = scorer.calculate_scores(comparison_bad, bad_data)
    feedback_bad = feedback_generator.generate_feedback(scores_bad, comparison_bad, bad_data)
    
    print(f"总分: {scores_bad['total']}/100")
    print(f"等级: {scorer.get_grade(scores_bad['total'])}")
    print(f"优点: {feedback_bad['strengths'][:2]}")
    print(f"不足: {feedback_bad['weaknesses'][:2]}")
    
    print("\n=== 测试报告导出 ===")
    report_data = {
        'original_filename': 'test_report.docx',
        'timestamp': '2024-01-01 12:00:00',
        'scores': scores_good,
        'grade': scorer.get_grade(scores_good['total']),
        'comparison': comparison_good,
        'feedback': feedback_good
    }
    
    json_path = exporter.export_json(report_data, 'test_report')
    html_path = exporter.export_html(report_data, 'test_report')
    
    print(f"JSON报告: {json_path}")
    print(f"HTML报告: {html_path}")
    
    print("\n=== 系统测试完成 ===")
    print("所有模块运行正常！")

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='测试实训报告智能批改系统')
    parser.add_argument('-c', '--create', action='store_true', help='仅创建测试数据')
    args = parser.parse_args()
    
    if args.create:
        create_test_data()
    else:
        test_system()