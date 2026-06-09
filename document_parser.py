import os
import re
from docx import Document
from docx.shared import Pt
import pdfplumber
from PIL import Image
from io import BytesIO

try:
    import win32com.client
    HAS_PYWIN32 = True
except ImportError:
    HAS_PYWIN32 = False

class DocumentParser:
    def __init__(self):
        self.supported_extensions = ['.docx', '.pdf', '.txt', '.md', '.rtf', '.odt', '.doc']
    
    def parse_document(self, file_path):
        """解析文档文件，返回文档内容和图片信息"""
        _, ext = os.path.splitext(file_path)
        
        if ext.lower() == '.docx':
            return self._parse_docx(file_path)
        elif ext.lower() == '.doc':
            return self._parse_doc(file_path)
        elif ext.lower() == '.pdf':
            return self._parse_pdf(file_path)
        elif ext.lower() == '.txt':
            return self._parse_txt(file_path)
        elif ext.lower() == '.md':
            return self._parse_md(file_path)
        elif ext.lower() == '.rtf':
            return self._parse_rtf(file_path)
        elif ext.lower() == '.odt':
            return self._parse_odt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def _parse_doc(self, file_path):
        """解析doc文件（旧版Word格式）"""
        content = []
        
        if HAS_PYWIN32:
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(file_path))
                
                for paragraph in doc.Paragraphs:
                    text = paragraph.Range.Text.strip()
                    if text:
                        content.append({
                            'type': 'paragraph',
                            'text': text,
                            'style': paragraph.Style.NameLocal if paragraph.Style else "Normal",
                            'is_heading': self._is_heading(text)
                        })
                
                doc.Close()
                word.Quit()
                
            except Exception as e:
                return self._parse_doc_fallback(file_path)
        else:
            return self._parse_doc_fallback(file_path)
        
        return {
            'content': content,
            'images': [],
            'tables': [],
            'total_text': ' '.join([item['text'] for item in content]),
            'image_count': 0
        }
    
    def _parse_doc_fallback(self, file_path):
        """doc文件解析失败时的备用方法"""
        content = []
        try:
            with open(file_path, 'rb') as f:
                raw = f.read()
                text = raw.decode('utf-8', errors='ignore')
                text = re.sub(r'[\x00-\x1f\x7f-\xff]', '', text)
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line and len(line) > 1:
                        content.append({
                            'type': 'paragraph',
                            'text': line,
                            'style': 'Normal',
                            'is_heading': self._is_heading(line)
                        })
        except:
            pass
        
        return {
            'content': content,
            'images': [],
            'tables': [],
            'total_text': ' '.join([item['text'] for item in content]),
            'image_count': 0
        }
    
    def _parse_txt(self, file_path):
        """解析txt文件"""
        content = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                for line in lines:
                    line = line.strip()
                    if line:
                        content.append({
                            'type': 'paragraph',
                            'text': line,
                            'style': 'Normal',
                            'is_heading': self._is_heading(line)
                        })
        except:
            with open(file_path, 'r', encoding='gbk') as f:
                lines = f.readlines()
                for line in lines:
                    line = line.strip()
                    if line:
                        content.append({
                            'type': 'paragraph',
                            'text': line,
                            'style': 'Normal',
                            'is_heading': self._is_heading(line)
                        })
        
        return {
            'content': content,
            'images': [],
            'tables': [],
            'total_text': ' '.join([item['text'] for item in content]),
            'image_count': 0
        }
    
    def _parse_md(self, file_path):
        """解析markdown文件"""
        content = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                for line in lines:
                    line = line.strip()
                    if line:
                        is_heading = False
                        style = 'Normal'
                        if line.startswith('#'):
                            is_heading = True
                            level = min(line.count('#'), 6)
                            style = f'Heading {level}'
                            line = line.lstrip('#').strip()
                        
                        content.append({
                            'type': 'paragraph',
                            'text': line,
                            'style': style,
                            'is_heading': is_heading
                        })
        except:
            with open(file_path, 'r', encoding='gbk') as f:
                lines = f.readlines()
                for line in lines:
                    line = line.strip()
                    if line:
                        is_heading = False
                        style = 'Normal'
                        if line.startswith('#'):
                            is_heading = True
                            level = min(line.count('#'), 6)
                            style = f'Heading {level}'
                            line = line.lstrip('#').strip()
                        
                        content.append({
                            'type': 'paragraph',
                            'text': line,
                            'style': style,
                            'is_heading': is_heading
                        })
        
        return {
            'content': content,
            'images': [],
            'tables': [],
            'total_text': ' '.join([item['text'] for item in content]),
            'image_count': 0
        }
    
    def _parse_rtf(self, file_path):
        """解析RTF文件"""
        content = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                
            text = re.sub(r'\\[a-z]+(\s+\d+)?', '', text)
            text = re.sub(r'{[^}]*}', '', text)
            text = re.sub(r'\\.', '', text)
            
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    content.append({
                        'type': 'paragraph',
                        'text': line,
                        'style': 'Normal',
                        'is_heading': self._is_heading(line)
                    })
        except:
            pass
        
        return {
            'content': content,
            'images': [],
            'tables': [],
            'total_text': ' '.join([item['text'] for item in content]),
            'image_count': 0
        }
    
    def _parse_odt(self, file_path):
        """解析ODT文件（简化版）"""
        content = []
        try:
            import zipfile
            from xml.etree import ElementTree
            
            with zipfile.ZipFile(file_path, 'r') as zf:
                with zf.open('content.xml') as f:
                    tree = ElementTree.parse(f)
                    namespace = {'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'}
                    
                    for elem in tree.iter():
                        if 'p' in elem.tag:
                            text = elem.text or ''
                            text = text.strip()
                            if text:
                                content.append({
                                    'type': 'paragraph',
                                    'text': text,
                                    'style': 'Normal',
                                    'is_heading': self._is_heading(text)
                                })
        except:
            pass
        
        if not content:
            try:
                with open(file_path, 'rb') as f:
                    raw = f.read()
                    text = raw.decode('utf-8', errors='ignore')
                    text = re.sub(r'<[^>]+>', '', text)
                    lines = text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            content.append({
                                'type': 'paragraph',
                                'text': line,
                                'style': 'Normal',
                                'is_heading': self._is_heading(line)
                            })
            except:
                pass
        
        return {
            'content': content,
            'images': [],
            'tables': [],
            'total_text': ' '.join([item['text'] for item in content]),
            'image_count': 0
        }
    
    def _parse_docx(self, file_path):
        """解析docx文件"""
        doc = Document(file_path)
        content = []
        images = []
        tables = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                paragraph = doc.paragraphs[len(content)]
                text = paragraph.text.strip()
                if text:
                    style = paragraph.style.name if paragraph.style else "Normal"
                    content.append({
                        'type': 'paragraph',
                        'text': text,
                        'style': style,
                        'is_heading': style.startswith('Heading')
                    })
            
            elif element.tag.endswith('tbl'):
                table_data = []
                table = doc.tables[len(tables)]
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                try:
                    img = Image.open(BytesIO(image_data))
                    images.append({
                        'width': img.width,
                        'height': img.height,
                        'size': len(image_data),
                        'format': img.format
                    })
                except:
                    images.append({
                        'width': 0,
                        'height': 0,
                        'size': len(image_data),
                        'format': 'unknown'
                    })
        
        return {
            'content': content,
            'images': images,
            'tables': tables,
            'total_text': ' '.join([item['text'] for item in content]),
            'image_count': len(images)
        }
    
    def _parse_pdf(self, file_path):
        """解析pdf文件"""
        content = []
        images = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        lines = text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                content.append({
                                    'type': 'paragraph',
                                    'text': line,
                                    'style': 'Normal',
                                    'is_heading': self._is_heading(line)
                                })
                    
                    for img in page.images:
                        try:
                            img_data = img['stream'].get_data()
                            img_obj = Image.open(BytesIO(img_data))
                            images.append({
                                'width': img_obj.width,
                                'height': img_obj.height,
                                'size': len(img_data),
                                'format': img_obj.format
                            })
                        except:
                            images.append({
                                'width': 0,
                                'height': 0,
                                'size': len(img_data),
                                'format': 'unknown'
                            })
        
        except Exception as e:
            raise ValueError(f"PDF解析失败: {str(e)}")
        
        return {
            'content': content,
            'images': images,
            'tables': [],
            'total_text': ' '.join([item['text'] for item in content]),
            'image_count': len(images)
        }
    
    def _is_heading(self, text):
        """简单判断是否为标题"""
        if len(text) < 30 and text.isupper():
            return True
        if re.match(r'^[一二三四五六七八九十]+[、.．]', text):
            return True
        if re.match(r'^[0-9]+[、.．]', text):
            return True
        if re.match(r'^(第[一二三四五六七八九十]+章|第[0-9]+章)', text):
            return True
        return False
    
    def extract_sections(self, doc_data, section_keywords):
        """从文档中提取章节"""
        sections = {}
        current_section = None
        
        for item in doc_data['content']:
            text = item['text']
            
            for keyword in section_keywords:
                if keyword in text or text in keyword:
                    current_section = keyword
                    sections[current_section] = []
                    break
            
            if current_section and item['text']:
                sections[current_section].append(item['text'])
        
        return sections